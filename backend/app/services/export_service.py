import docx
import openpyxl
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import io
from typing import List, Dict, Any

class ExportService:
    def export_docx(self, patent_num: str, ref_num: str, rows: List[Dict[str, Any]]) -> io.BytesIO:
        doc = docx.Document()
        doc.add_heading(f"Patent Invalidity Obviousness Chart", level=0)
        doc.add_paragraph(f"Subject Patent: {patent_num} vs Reference Prior Art: {ref_num}")
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Element ID'
        hdr_cells[1].text = 'Claim limitation (Subject)'
        hdr_cells[2].text = 'Prior Art Citation'
        hdr_cells[3].text = 'Location'
        hdr_cells[4].text = 'Analysis & Rationale'
        
        for r in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(r.get("element_id", ""))
            row_cells[1].text = str(r.get("claim_text", ""))
            row_cells[2].text = str(r.get("cited_passage", ""))
            row_cells[3].text = f"Para: {r.get('para_ref','')}\nFig: {r.get('fig_ref','')}"
            row_cells[4].text = str(r.get("rationale", ""))
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def export_xlsx(self, patent_num: str, ref_num: str, rows: List[Dict[str, Any]]) -> io.BytesIO:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Claim Chart"
        
        ws.append([f"Patent Invalidity Analysis Chart"])
        ws.append([f"Subject Patent: {patent_num}", f"Reference Art: {ref_num}"])
        ws.append([]) # empty separator row
        
        ws.append(['Element ID', 'Claim limitation (Subject)', 'Prior Art Citation', 'Location References', 'Obviousness Rationale'])
        for r in rows:
            ws.append([
                r.get("element_id", ""),
                r.get("claim_text", ""),
                r.get("cited_passage", ""),
                f"Para: {r.get('para_ref','')}, Fig: {r.get('fig_ref','')}",
                r.get("rationale", "")
            ])
            
        # Adjust column widths slightly
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 25
            
        file_stream = io.BytesIO()
        wb.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def export_pdf(self, patent_num: str, ref_num: str, rows: List[Dict[str, Any]]) -> io.BytesIO:
        file_stream = io.BytesIO()
        doc = SimpleDocTemplate(
            file_stream,
            pagesize=landscape(A4),
            rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30
        )
        
        styles = getSampleStyleSheet()
        normal = styles['Normal']
        
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1E3A8A'),
            spaceAfter=15
        )
        
        story = [
            Paragraph(f"Patent Invalidity Obviousness Analysis Chart", title_style),
            Paragraph(f"Subject Patent: {patent_num} | Prior Art Reference: {ref_num}", normal),
            Spacer(1, 15)
        ]
        
        data = [['ID', 'Claim limitation (Subject)', 'Prior Art Citation', 'Location (Ref)', 'Analysis & Rationale']]
        for r in rows:
            data.append([
                Paragraph(str(r.get("element_id", "")), normal),
                Paragraph(str(r.get("claim_text", "")), normal),
                Paragraph(str(r.get("cited_passage", "")), normal),
                Paragraph(f"Para: {r.get('para_ref','')}<br/>Fig: {r.get('fig_ref','')}", normal),
                Paragraph(str(r.get("rationale", "")), normal)
            ])
            
        table = Table(data, colWidths=[40, 180, 180, 80, 260])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E3A8A')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,0), 8),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#F3F4F6'), colors.white]),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#D1D5DB')),
        ]))
        
        story.append(table)
        doc.build(story)
        file_stream.seek(0)
        return file_stream

export_service = ExportService()
